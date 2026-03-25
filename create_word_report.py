"""
Create Professional Word Document for SHELDON Capabilities Report
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('SHELDON Executive Interface', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Comprehensive Capabilities Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

# Date and purpose
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n').bold = True
info.add_run('Prepared for: Dennis Straub, President & CEO\n')
info.add_run('AmeriQual Foods')

doc.add_paragraph()

# Executive Summary Box
doc.add_heading('Executive Summary', level=1)
summary = doc.add_paragraph()
summary.add_run('SHELDON is an executive intelligence dashboard that provides real-time visibility into AmeriQual Foods operations, finances, and inventory. ')
summary.add_run('The system is approximately 80% complete').bold = True
summary.add_run(', with live data connections to Snowflake (production/OEE), Sage X3 (financials/inventory), and OpenAI (AI chat assistant).')

doc.add_paragraph()

# Current Capabilities Section
doc.add_heading('1. Current Capabilities', level=1)

doc.add_heading('Data Sources Connected', level=2)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Source'
hdr_cells[1].text = 'Type'
hdr_cells[2].text = 'Status'
hdr_cells[3].text = 'Data Provided'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True
    shading = cell._element.get_or_add_tcPr()

# Data rows
data = [
    ('Snowflake', 'Redzone MES', 'Active', 'OEE, production, downtime, labor'),
    ('Sage X3', 'SQL Server', 'Active', 'Revenue, inventory, AR, financials'),
    ('OpenAI', 'AI Chat', 'Active', 'gpt-4o-mini for SHELDON persona'),
]
for i, row_data in enumerate(data):
    row = table.rows[i + 1].cells
    for j, text in enumerate(row_data):
        row[j].text = text

doc.add_paragraph()

doc.add_heading('Live KPIs Currently Displayed', level=2)

doc.add_heading('Operations (from Snowflake/Redzone)', level=3)
ops_kpis = [
    'Plant OEE (Overall Equipment Effectiveness)',
    'Quality, Availability, Performance metrics',
    'Total Output (24-hour production units)',
    'Downtime Hours (planned + unplanned)',
    'Lines Reporting / Active Lines',
    'Problem Lines (below 70% OEE threshold)',
]
for kpi in ops_kpis:
    doc.add_paragraph(kpi, style='List Bullet')

doc.add_heading('Financial (from Sage X3)', level=3)
fin_kpis = [
    'Revenue: Daily / MTD / YTD',
    'Inventory Total Value: $108.4M (FD1: $44.7M, PK1: $16.1M, TP1: $47.6M)',
    'AR Days: 45 days',
    'Cash Position: $3.0M',
    'Gross Margin: 19.4%',
    'Inventory Turns: 111.0x',
]
for kpi in fin_kpis:
    doc.add_paragraph(kpi, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Features Available', level=2)
table2 = doc.add_table(rows=11, cols=3)
table2.style = 'Table Grid'

hdr = table2.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'Status'
hdr[2].text = 'Description'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

features = [
    ('Dashboard', 'Working', '8 KPI cards with comparison badges'),
    ('AI Chat', 'Working', 'OpenAI-powered SHELDON persona'),
    ('Text-to-Speech', 'Working', 'Browser-native, zero cost'),
    ('Executive Briefing', 'Working', 'Auto-generated summary on load'),
    ('Red Flags Panel', 'Partial', 'Framework ready, needs threshold config'),
    ('PDF Export', 'Working', 'Full dashboard screenshot'),
    ('Data Polling', 'Working', '60-second auto-refresh'),
    ('Operations Panel', 'Working', 'Downtime breakdown, OEE by line'),
    ('Financial Panel', 'Working', 'Live revenue, margin, cash'),
    ('Inventory Panel', 'Working', 'Value by facility, FG breakdown'),
]
for i, row_data in enumerate(features):
    row = table2.rows[i + 1].cells
    for j, text in enumerate(row_data):
        row[j].text = text

doc.add_page_break()

# Current Limitations Section
doc.add_heading('2. Current Limitations & Gaps', level=1)

doc.add_heading('Incomplete Panels', level=2)
table3 = doc.add_table(rows=7, cols=3)
table3.style = 'Table Grid'

hdr = table3.rows[0].cells
hdr[0].text = 'Panel'
hdr[1].text = 'Status'
hdr[2].text = 'Blocker'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

gaps = [
    ('People', 'Framework only', 'No HR/payroll data source'),
    ('Quality', 'Framework only', 'No quality system integration'),
    ('Marketing', 'Framework only', 'No CRM/pipeline data'),
    ('INST Business', 'Framework only', 'Needs INST-specific data split'),
    ('Board Report', 'Placeholder', 'Not implemented'),
    ('What-If Analysis', 'Placeholder', 'Concept only'),
]
for i, row_data in enumerate(gaps):
    row = table3.rows[i + 1].cells
    for j, text in enumerate(row_data):
        row[j].text = text

doc.add_paragraph()

doc.add_heading('Missing Metrics', level=2)
missing = [
    'EBITDA (optional - Gross Margin covers profitability)',
    'Schedule Attainment (exists in Redzone, not yet queried)',
    'Period comparisons (vs last week, vs budget)',
    'Inventory threshold alerts (SKUs below safety stock)',
    'Quality holds and rework tracking',
    'Customer profitability analysis',
    'Cash flow forecasting',
]
for item in missing:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Known Technical Issues', level=2)
issues = [
    'Red flags query currently failing (SnowSQL configuration) - fixed to fail gracefully',
    'Briefing load time: 55-60 seconds (needs caching strategy)',
    'No conversation persistence (chat resets on browser refresh)',
    'Dark theme may be unprofessional for board presentations',
]
for item in issues:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Questions for Dennis Section
doc.add_heading('3. Questions for Dennis', level=1)
intro = doc.add_paragraph()
intro.add_run('The following questions will help optimize SHELDON for executive decision-making. ').italic = True

doc.add_heading('A. KPI Priorities & Thresholds', level=2)
questions_a = [
    'Which 3 KPIs matter most for your morning glance? (Currently: OEE, Revenue, Inventory)',
    'What is your OEE alert threshold? (Currently: <70% triggers red flag)',
    'What is your target Gross Margin? (Currently showing: 19.4%)',
    'How many AR Days are acceptable? (Currently: 45 days)',
    'What is your target Schedule Attainment? (Not currently displayed)',
    'Is $108M the right inventory level? What is the target turnover rate?',
]
for i, q in enumerate(questions_a, 1):
    doc.add_paragraph(f'{i}. {q}')

doc.add_heading('B. Workflow & Usage Patterns', level=2)
questions_b = [
    'When do you check SHELDON? (Morning briefing? Throughout day? Before meetings?)',
    'Should voice briefing auto-play, or be on-demand only?',
    'Who else will use SHELDON? (Other executives? Operations team? Board members?)',
    'What action do you take when you see a red flag? Who do you call?',
]
for i, q in enumerate(questions_b, 7):
    doc.add_paragraph(f'{i}. {q}')

doc.add_heading('C. Missing Data Sources', level=2)
questions_c = [
    'Schedule Attainment: How is it calculated? Where is it tracked?',
    'Quality data: Where is it stored? What metrics matter most?',
    'People/Labor insights: Want staffing levels? Overtime? Turnover?',
    'Customer profitability: Which customers drive most profit?',
    'Cash position: How often is it updated? What is the minimum threshold?',
]
for i, q in enumerate(questions_c, 11):
    doc.add_paragraph(f'{i}. {q}')

doc.add_heading('D. Advanced Features', level=2)
questions_d = [
    'Do you want Forecast vs Actual comparisons (Budget vs Actual revenue)?',
    'Would you like AI-categorized downtime analysis? (We have 32,338 operator comments)',
    'Should facilities be ranked by Revenue? OEE? Margin?',
    'Would you use What-If scenarios? ("What if we run one less shift?")',
    'Should SHELDON auto-generate board summaries? How often?',
]
for i, q in enumerate(questions_d, 16):
    doc.add_paragraph(f'{i}. {q}')

doc.add_heading('E. Design & Presentation', level=2)
questions_e = [
    'Is the current dark theme appropriate for board meetings? Need a light theme?',
    'Is PDF export sufficient? Need Excel or PowerPoint export?',
    'Do you need mobile access to check SHELDON while traveling?',
    'Is 60-second data refresh acceptable? Need real-time?',
    'Should the chat remember previous conversations across sessions?',
]
for i, q in enumerate(questions_e, 21):
    doc.add_paragraph(f'{i}. {q}')

doc.add_page_break()

# Strategic Summary
doc.add_heading('4. Strategic Summary', level=1)

doc.add_heading('Current State: 80% Complete', level=2)

doc.add_heading('Strengths', level=3)
strengths = [
    'Direct database integration (no external service dependency for data)',
    'Live financial + operational data combined in single dashboard',
    'Comprehensive data available (91 Snowflake views, 1,400+ Sage tables)',
    '32,338 operator downtime comments available for AI analysis',
]
for item in strengths:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Critical Gaps', level=3)
critical = [
    'Red flags need threshold configuration and ownership assignment',
    'Metrics lack historical comparison (vs yesterday, vs budget, vs target)',
    '6 of 7 category panels still placeholder',
    'Briefing takes 55-60 seconds to load (needs optimization)',
]
for item in critical:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Recommended Next Steps', level=2)
steps = [
    ('1.', 'Meet with Dennis', 'Define red flag thresholds and KPI priorities'),
    ('2.', 'Add comparison context', '"vs yesterday", "vs target" on all KPI cards'),
    ('3.', 'Build Schedule Attainment', 'High-value missing metric from Redzone'),
    ('4.', 'Optimize briefing performance', 'Implement caching strategy'),
    ('5.', 'Consider light theme option', 'For board presentations'),
]

table4 = doc.add_table(rows=6, cols=3)
table4.style = 'Table Grid'
hdr = table4.rows[0].cells
hdr[0].text = 'Priority'
hdr[1].text = 'Action'
hdr[2].text = 'Description'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

for i, row_data in enumerate(steps):
    row = table4.rows[i + 1].cells
    for j, text in enumerate(row_data):
        row[j].text = text

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('Report generated from comprehensive analysis of all SHELDON project files\n').italic = True
footer.add_run('AmeriQual Foods - Confidential').bold = True

# Save document
output_path = r"C:\Users\Claude\Business\AMERIQUAL PROJECT TRACKER\Current Projects\SHELDON Executive Interface\SHELDON Capabilities Report.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
